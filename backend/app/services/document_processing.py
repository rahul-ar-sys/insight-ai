import os
import asyncio
import aiofiles
from typing import List, Optional, Dict, Any
from uuid import UUID
from sqlalchemy.orm import Session
from sentence_transformers import SentenceTransformer
import pandas as pd
import numpy as np
from PIL import Image
import pytesseract
import PyPDF2
from docx import Document as DocxDocument
import io
import tempfile
import hashlib
from datetime import datetime
import concurrent.futures
import torch
import time
import json
from pdf2image import convert_from_bytes 
from ..core.logging import logger, log_agent_trace

from ..core.config import settings
from ..core.database import get_redis, SessionLocal
from ..models.database import Document, DocumentChunk
from ..models.document_status import DocumentStatus, ProcessingStep, ProcessingStepStatus
from .storage import StorageService
from .vector_store import VectorStoreService

from ..dependencies import get_document_processing_service, get_vector_store_service


class DocumentProcessingService:
    """Service for processing uploaded documents"""
    
    def __init__(self, vector_service: VectorStoreService):
        logger.info("Initializing DocumentProcessingService...")
        self.storage_service = StorageService()
        self.vector_service = vector_service
        self.redis = None
        self.embedding_model = None
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.executor = concurrent.futures.ThreadPoolExecutor(max_workers=os.cpu_count())

    async def _get_redis_client(self):
        """Initializes the redis client if it's not already."""
        if self.redis is None:
            self.redis = await get_redis()
        return self.redis

    async def load_models(self):
        """Asynchronously loads the embedding model. To be called at application startup."""
        if self.embedding_model:
            logger.info("Embedding model already loaded.")
            return
        await asyncio.to_thread(self._load_embedding_model_sync)

    def _load_embedding_model_sync(self):
        """Load sentence transformer model for embeddings with optimizations"""
        try:
            model_name = settings.embedding_model
            print(f"DEBUG: Attempting to load model from path: '{model_name}'")
            self.embedding_model = SentenceTransformer(model_name, device=self.device)
            logger.info(f"Loaded embedding model on {self.device}: {model_name}")
        except Exception as e:
            logger.error(f"Failed to load embedding model: {e}", exc_info=True)
            self.embedding_model = None

    async def _update_document_status(self, document_id: UUID, status: DocumentStatus):
        """Saves the document's processing status to Redis."""
        try:
            redis_client = await self._get_redis_client()
            status_key = f"document_status:{document_id}"
            await redis_client.set(status_key, status.json(), ex=3600)
        except Exception as e:
            logger.error(f"Failed to update document status in Redis for {document_id}: {e}")

    async def process_document_from_temp(self, document_id: UUID, temp_file_path: str, db_session_factory):
        """
        Reads a document from a temporary file path, stores it permanently,
        and then processes it. Cleans up the temporary file afterward.
        """
        db = next(db_session_factory())
        document = db.query(Document).filter(Document.document_id == document_id).first()
        if not document:
            logger.error(f"Document {document_id} not found for processing from temp file.")
            db.close()
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
            return

        try:
            async with aiofiles.open(temp_file_path, 'rb') as f:
                file_content = await f.read()

            storage_path = await self.storage_service.store_file(
                file_content,
                f"{document.workspace_id}/{document.document_id}_{document.original_name}"
            )

            document.storage_path = storage_path
            db.commit()
            db.refresh(document)

            await self.process_document(document_id, file_content, db_session_factory)

        finally:
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
                logger.info(f"Cleaned up temporary file: {temp_file_path}")

    async def process_document(self, document_id: UUID, file_content: bytes, db_session_factory):
        """Process uploaded document through the optimized pipeline"""
        
        db = next(db_session_factory())
        document = db.query(Document).filter(Document.document_id == document_id).first()
        if not document:
            logger.error(f"Document {document_id} not found in database for background processing.")
            db.close()
            return

        status = DocumentStatus(
            document_id=document.document_id,
            file_name=document.original_name,
            status="processing",
            progress=0,
            steps=[
                ProcessingStep(name="File Storage & Text Extraction", status=ProcessingStepStatus.PENDING),
                ProcessingStep(name="Semantic Chunking", status=ProcessingStepStatus.PENDING),
                ProcessingStep(name="Embedding Generation", status=ProcessingStepStatus.PENDING),
                ProcessingStep(name="Database Indexing", status=ProcessingStepStatus.PENDING)
            ]
        )
        await self._update_document_status(document.document_id, status)

        processing_start_time = time.time()
        
        try:
            log_agent_trace("document_processor", "start_processing", {"document_id": str(document.document_id), "file_type": document.file_type})
            logger.info(f"🚀 Starting optimized processing for document {document.document_id}")

            status.steps[0].status = ProcessingStepStatus.IN_PROGRESS
            await self._update_document_status(document.document_id, status)
            
            text_content = await self._extract_text_optimized(file_content, document.file_type)
            
            if not text_content:
                error_msg = f"Could not extract any text from the document '{document.original_name}'. The file might be corrupted, password-protected, or an image-only document (e.g., a scanned PDF)."
                if document.file_type == "application/pdf":
                    error_msg += " For scanned PDFs, an OCR tool is required to extract text, which is not enabled by default."
                raise ValueError(error_msg)
            
            status.steps[0].status = ProcessingStepStatus.COMPLETED
            status.progress = 25
            await self._update_document_status(document.document_id, status)
            logger.info(f"📖 Text extraction completed: {len(text_content)} characters")
            
            status.steps[1].status = ProcessingStepStatus.IN_PROGRESS
            await self._update_document_status(document.document_id, status)

            chunks = await self._create_chunks_optimized(text_content)
            
            status.steps[1].status = ProcessingStepStatus.COMPLETED
            status.progress = 50
            await self._update_document_status(document.document_id, status)
            logger.info(f"✂️ Created {len(chunks)} optimized chunks")
            
            status.steps[2].status = ProcessingStepStatus.IN_PROGRESS
            status.steps[3].status = ProcessingStepStatus.IN_PROGRESS
            await self._update_document_status(document.document_id, status)

            logger.info(f"-> Starting batch processing for {len(chunks)} chunks...")
            await self._process_chunks_batch(chunks, document, db)
            logger.info("-> Finished batch processing successfully.")
            
            status.steps[2].status = ProcessingStepStatus.COMPLETED
            status.progress = 75
            await self._update_document_status(document.document_id, status)
            
            status.steps[3].status = ProcessingStepStatus.COMPLETED
            status.progress = 99
            await self._update_document_status(document.document_id, status)

            document.status = "ready"
            document.total_chunks = len(chunks)
            document.embeddings_generated = True
            db.commit()
            
            processing_time = time.time() - processing_start_time
            log_agent_trace("document_processor", "processing_complete", {"document_id": str(document.document_id), "chunks_created": len(chunks), "ocr_applied": True, "processing_time": processing_time})
            logger.info(f"✅ Successfully processed document {document.document_id} in {processing_time:.2f}s")
            
        except Exception as e:
            processing_time = time.time() - processing_start_time
            logger.error(f"❌ Error processing document {document.document_id} after {processing_time:.2f}s: {e}", exc_info=True)
            
            document.status = "error"
            document.error_message = str(e)
            db.commit()
            
            status.status = "error"
            status.error_message = str(e)
            for step in status.steps:
                if step.status == ProcessingStepStatus.IN_PROGRESS:
                    step.status = ProcessingStepStatus.FAILED
            await self._update_document_status(document.document_id, status)

        finally:
            status.progress = 100
            if status.status != "error":
                status.status = "ready"
            await self._update_document_status(document.document_id, status)
            db.close()

    async def _extract_text_optimized(self, file_content: bytes, file_type: str) -> str:
        loop = asyncio.get_event_loop()
        
        def extract_text():
            if file_type == "application/pdf":
                return self._extract_pdf_optimized(file_content)
            elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                return self._extract_docx_optimized(file_content)
            elif file_type.startswith("text/"):
                return file_content.decode('utf-8', errors='ignore')
            else:
                logger.warning(f"Unsupported file type for text extraction: {file_type}")
                return ""
        
        return await loop.run_in_executor(self.executor, extract_text)
    
    # --- REPLACE THIS ENTIRE FUNCTION ---
    def _extract_pdf_optimized(self, file_content: bytes) -> str:
        # --- First, try standard text extraction ---
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            with concurrent.futures.ThreadPoolExecutor() as executor:
                page_futures = [executor.submit(self._extract_page_text, page, i) for i, page in enumerate(pdf_reader.pages)]
                pages_text_tuples = [future.result() for future in concurrent.futures.as_completed(page_futures)]
            
            pages_text_tuples.sort() # Sort by page number
            text_content = "\n\n".join([text for _, text in pages_text_tuples if text])
            
            if text_content.strip():
                logger.info("Successfully extracted digital text from PDF.")
                return text_content
        except Exception as e:
            logger.warning(f"PyPDF2 failed to extract text, will attempt OCR. Reason: {e}")

        # --- If standard extraction fails or yields no text, perform OCR ---
        logger.info("No digital text found. Attempting OCR with Tesseract...")
        try:
            # --- ADD THIS DEBUG LINE ---
            print(f"DEBUG: Application is using Tesseract path: '{settings.tesseract_path}'")

            if settings.tesseract_path and os.path.exists(settings.tesseract_path):
                pytesseract.pytesseract.tesseract_cmd = settings.tesseract_path

            images = convert_from_bytes(file_content)
            ocr_texts = []
            for i, image in enumerate(images):
                page_text = pytesseract.image_to_string(image)
                ocr_texts.append(f"--- Page {i + 1} ---\n\n{page_text}")
            
            full_text = "\n\n".join(ocr_texts)
            logger.info(f"Successfully extracted text via OCR from {len(images)} pages.")
            return full_text
        except Exception as ocr_error:
            logger.error(f"OCR extraction failed: {ocr_error}", exc_info=True)
            return ""

    def _extract_page_text(self, page, page_num: int) -> tuple:
        try:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                return page_num, f"--- Page {page_num + 1} ---\n\n{page_text}"
            return page_num, ""
        except Exception as e:
            logger.warning(f"Error extracting page {page_num}: {e}")
            return page_num, ""
    
    def _extract_docx_optimized(self, file_content: bytes) -> str:
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            table_texts = ["\t".join(cell.text for cell in row.cells) for table in doc.tables for row in table.rows]
            return "\n".join(paragraphs) + "\n\n" + "\n".join(filter(None, table_texts))
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""

    async def _create_chunks_optimized(self, text: str) -> List[Dict[str, Any]]:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            is_separator_regex=False,
        )
        
        split_texts = text_splitter.split_text(text)
        
        chunks = [{
            "chunk_index": i,
            "content": chunk_text,
        } for i, chunk_text in enumerate(split_texts)]
        
        logger.info(f"📄 Created {len(chunks)} optimized chunks")
        return chunks

    async def _process_chunks_batch(self, chunks: List[Dict[str, Any]], document: Document, db: Session):
        if not self.embedding_model:
            raise Exception("Embedding model not available")
        
        chunk_texts = [chunk["content"] for chunk in chunks]
        
        try:
            batch_size = 32
            all_embeddings = []
            logger.info("--> Starting embedding generation...")
            for i in range(0, len(chunk_texts), batch_size):
                batch_texts = chunk_texts[i:i + batch_size]
                batch_embeddings = self.embedding_model.encode(
                    batch_texts,
                    batch_size=len(batch_texts),
                    show_progress_bar=False,
                    convert_to_numpy=True,
                    normalize_embeddings=True
                )
                all_embeddings.extend(batch_embeddings)
            logger.info(f"--> Finished embedding generation. Total embeddings: {len(all_embeddings)}")
        except Exception as e:
            logger.error(f"❌ Batch embedding generation failed: {e}", exc_info=True)
            raise

        logger.info("--> Starting vector storage...")
        storage_tasks = []
        for i, chunk in enumerate(chunks):
            metadata = {
                "document_id": str(document.document_id),
                "workspace_id": str(document.workspace_id),
                "chunk_index": chunk["chunk_index"],
                "file_name": document.original_name,
            }
            storage_tasks.append(
                self.vector_service.store_embedding(
                    text=chunk["content"],
                    embedding=all_embeddings[i],
                    metadata=metadata
                )
            )
        
        vector_ids = await asyncio.gather(*storage_tasks)
        logger.info(f"--> Finished vector storage. Total vectors stored: {len(vector_ids)}")

        db_chunks = []
        for i, chunk_data in enumerate(chunks):
            db_chunk = DocumentChunk(
                document_id=document.document_id,
                chunk_index=chunk_data["chunk_index"],
                content=chunk_data["content"],
                vector_id=vector_ids[i]
            )
            db_chunks.append(db_chunk)
        
        db.add_all(db_chunks)
        db.commit()